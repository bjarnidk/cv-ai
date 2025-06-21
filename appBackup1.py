from dotenv import load_dotenv
load_dotenv()

import streamlit as st
import PyPDF2
import os
from io import BytesIO
from openai import OpenAI
from docx import Document

# Theme and layout
st.set_page_config(page_title="Zempli", layout="centered")

# Custom styling to match Carrd
st.markdown("""
    <style>
        div.stButton > button {
            background-color: #FDE68A;
            color: black;
            font-weight: bold;
            border: 1px solid #00000033;
            border-radius: 12px;
            padding: 0.6em 2em;
        }
    </style>
""", unsafe_allow_html=True)

# Custom header
st.markdown("""
    <div style='text-align: center; margin-top: 2rem;'>
        <div style='display: inline-block; background-color: #A7F3D0; padding: 8px 20px; border-radius: 12px; font-weight: bold; font-size: 28px; color: black;'>Zempli</div>
        <div style='margin-top: 1.5rem; background-color: #C4B5FD; padding: 18px 24px; border-radius: 16px; display: inline-block; font-family: monospace; font-size: 16px; color: #111827; max-width: 600px;'>
            Upload your <b>CV</b> and a <b>job posting</b>.<br>
            Our AI will instantly generate a tailored, personal cover letter — in any language.
        </div>
    </div>
""", unsafe_allow_html=True)

# Initialize OpenAI client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Upload CV
cv_file = st.file_uploader("📄 Upload your CV (PDF or Word)", type=["pdf", "docx"])

# Tone selector
tone = st.selectbox("🎨 Choose tone for your application:", [
    "Professionel", "Varm og empatisk", "Selvsikker og ambitiøs", "Kortfattet og ligetil"
])

# Job description
job_description = st.text_area("📋 Paste the job description here:")

if st.button("✨ Try it for free"):
    if not cv_file or not job_description:
        st.warning("Please upload a CV and paste a job description.")
        st.stop()

    # Extract text from uploaded CV
    cv_text = ""
    if cv_file.name.endswith(".pdf"):
        pdf_reader = PyPDF2.PdfReader(cv_file)
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                cv_text += text + "\n"
    elif cv_file.name.endswith(".docx"):
        doc = Document(cv_file)
        for para in doc.paragraphs:
            cv_text += para.text + "\n"

    # Detect language
    with st.spinner("🔍 Detecting language..."):
        lang_response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a language expert."},
                {"role": "user", "content": f"What language is this CV written in?\n\n{cv_text[:1000]}"}
            ],
            temperature=0,
            max_tokens=20
        )
        detected_language = lang_response.choices[0].message.content.strip()

    # Generate cover letter
    with st.spinner("✍️ Generating tailored application..."):
        prompt_instruction = f"""
Write the following in {detected_language}:

Use a tone that is {tone.lower()}.

Use the CV and job description below to write a **tailored** and **personal** cover letter. 
Use real strengths, courses, and experiences from the CV that are relevant to this job.
Reference specific responsibilities or values from the job description where appropriate.

Here is the CV:
{cv_text}

Here is the job description:
{job_description}
"""

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an expert in writing job applications tailored to each candidate."},
                {"role": "user", "content": prompt_instruction}
            ],
            temperature=0.7,
            max_tokens=1000
        )

        cover_letter = response.choices[0].message.content

        st.subheader(f"📄 Din AI-genererede ansøgning ({detected_language}):")
        edited_text = st.text_area("📝 Edit the letter before download:", cover_letter, height=400)

        # Prepare Word file
        docx_file = BytesIO()
        doc = Document()
        doc.add_paragraph(edited_text)
        doc.save(docx_file)
        docx_file.seek(0)

        st.download_button("⬇️ Download as Word (.docx)", data=docx_file, file_name="cover_letter.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
